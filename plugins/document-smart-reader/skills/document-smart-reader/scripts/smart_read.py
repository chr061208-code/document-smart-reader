#!/usr/bin/env python3
"""Build and query a page-aware local cache for PDF and Word documents."""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path


SUPPORTED = {".pdf", ".docx"}
DEFAULT_CHUNK_CHARS = 6000
TOOL_VERSION = "0.2.0"

# Some PDFs contain incomplete font descriptors. pdfminer can still extract
# their text, but otherwise emits one warning per glyph/page and obscures the
# useful JSON result.
logging.getLogger("pdfminer").setLevel(logging.ERROR)


def normalize_text(value: str) -> str:
    """Normalize compatibility glyphs and remove extraction-only controls."""
    value = unicodedata.normalize("NFKC", value).replace("\u00a0", " ")
    value = "".join(char for char in value if char in "\n\t" or unicodedata.category(char) != "Cc")
    lines = [re.sub(r"[ \t]+", " ", line).rstrip() for line in value.splitlines()]
    return "\n".join(lines).strip()


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def default_cache_root() -> Path:
    override = os.environ.get("CODEX_DOCUMENT_READER_CACHE")
    if override:
        return Path(override).expanduser()
    codex_docs = Path.home() / "Documents" / "Codex"
    return (codex_docs if codex_docs.exists() else Path.cwd()) / ".document-reader-cache"


def find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return str(candidate)
    return None


def convert_docx_to_pdf(source: Path, cache_dir: Path) -> Path | None:
    soffice = find_soffice()
    if not soffice:
        return None
    rendered = cache_dir / "rendered"
    rendered.mkdir(parents=True, exist_ok=True)
    target = rendered / f"{source.stem}.pdf"
    if target.exists():
        return target
    profile_dir = Path(tempfile.mkdtemp(prefix="codex-lo-"))
    try:
        profile_uri = profile_dir.resolve().as_uri()
        result = subprocess.run(
            [soffice, f"-env:UserInstallation={profile_uri}", "--headless", "--convert-to", "pdf", "--outdir", str(rendered), str(source)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            check=False,
        )
        if result.returncode != 0 or not target.exists():
            return None
        return target
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


def extract_pdf(path: Path) -> tuple[list[dict], list[str]]:
    warnings: list[str] = []
    try:
        import pdfplumber  # type: ignore

        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for number, page in enumerate(pdf.pages, 1):
                text = normalize_text(page.extract_text(x_tolerance=2, y_tolerance=3) or "")
                words = page.extract_words() or []
                flags: list[str] = []
                if len(text) < 8:
                    flags.append("no_text")
                elif len(text) < 80:
                    flags.append("low_text")
                drawing_count = len(getattr(page, "lines", []) or []) + len(getattr(page, "rects", []) or [])
                image_count = len(getattr(page, "images", []) or [])
                if image_count or drawing_count >= 20:
                    flags.append("possible_complex_layout")
                pages.append({"page": number, "text": text, "flags": flags})
        return pages, warnings
    except Exception as exc:
        warnings.append(f"pdfplumber failed: {exc}")

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        pages = []
        for number, page in enumerate(reader.pages, 1):
            text = normalize_text(page.extract_text() or "")
            flags = ["no_text"] if len(text) < 8 else (["low_text"] if len(text) < 80 else [])
            pages.append({"page": number, "text": text, "flags": flags})
        return pages, warnings
    except Exception as exc:
        raise RuntimeError(f"Unable to extract PDF text: {exc}") from exc


def extract_docx_fallback(path: Path) -> tuple[list[dict], list[str]]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"DOCX fallback requires python-docx: {exc}") from exc

    document = Document(str(path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        blocks.append(f"## {text}" if "heading" in style else text)
    for table_number, table in enumerate(document.tables, 1):
        blocks.append(f"\n### Table {table_number}")
        for row in table.rows:
            blocks.append(" | ".join(normalize_text(cell.text).replace("\n", " ") for cell in row.cells))
    text = normalize_text("\n\n".join(blocks))
    return [{"page": None, "section": 1, "text": text, "flags": ["page_numbers_unavailable"]}], [
        "LibreOffice conversion unavailable; used structural DOCX extraction without reliable page numbers."
    ]


def page_markdown(page: dict) -> str:
    if page.get("page") is not None:
        anchor = f"<!-- source-page: {page['page']} -->\n\n## Page {page['page']}"
    else:
        anchor = f"<!-- source-section: {page.get('section', 1)} -->\n\n## Section {page.get('section', 1)}"
    text = normalize_text(page.get("text", "")) or "[No extractable text]"
    return f"{anchor}\n\n{text.strip()}\n"


def make_chunks(pages: list[dict], target_chars: int) -> list[dict]:
    chunks: list[dict] = []
    current: list[dict] = []
    current_chars = 0
    for page in pages:
        rendered = page_markdown(page)
        if current and current_chars + len(rendered) > target_chars:
            chunks.append(build_chunk(current, len(chunks) + 1))
            current, current_chars = [], 0
        current.append(page)
        current_chars += len(rendered)
    if current:
        chunks.append(build_chunk(current, len(chunks) + 1))
    return chunks


def build_chunk(pages: list[dict], number: int) -> dict:
    text = "\n\n".join(page_markdown(page) for page in pages)
    page_numbers = [page["page"] for page in pages if page.get("page") is not None]
    flags = sorted({flag for page in pages for flag in page.get("flags", [])})
    return {
        "id": number,
        "file": f"chunks/chunk-{number:04d}.md",
        "page_start": min(page_numbers) if page_numbers else None,
        "page_end": max(page_numbers) if page_numbers else None,
        "chars": len(text),
        "flags": flags,
        "text": text,
    }


def tokenize(text: str) -> list[str]:
    lowered = text.lower()
    tokens = re.findall(r"[a-z0-9_]+", lowered)
    for run in re.findall(r"[\u3400-\u9fff]+", lowered):
        tokens.extend(run if len(run) == 1 else (run[i : i + 2] for i in range(len(run) - 1)))
    return tokens


def build_index(chunks: list[dict]) -> list[dict]:
    records = []
    for chunk in chunks:
        counts = Counter(tokenize(chunk["text"]))
        records.append({
            "id": chunk["id"],
            "file": chunk["file"],
            "page_start": chunk["page_start"],
            "page_end": chunk["page_end"],
            "chars": chunk["chars"],
            "flags": chunk["flags"],
            "terms": dict(counts),
        })
    return records


def make_snippet(text: str, question: str, query_terms: list[str], width: int = 500) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    lowered = compact.lower()
    exact = question.lower().strip()
    if exact and exact in lowered:
        position = lowered.find(exact)
    else:
        candidates = sorted(
            (term for term in set(query_terms) if term and term in lowered),
            key=lambda term: (lowered.count(term), -len(term)),
        )
        position = lowered.find(candidates[0]) if candidates else 0
    start = max(0, position - width // 3)
    end = min(len(compact), start + width)
    prefix = "…" if start else ""
    suffix = "…" if end < len(compact) else ""
    return f"{prefix}{compact[start:end]}{suffix}"


def prepare(args: argparse.Namespace) -> int:
    source = Path(args.file).expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    if source.suffix.lower() not in SUPPORTED:
        raise ValueError(f"Unsupported format: {source.suffix}. Supported: {', '.join(sorted(SUPPORTED))}")

    digest = file_hash(source)
    cache_root = Path(args.output).expanduser().resolve() if args.output else default_cache_root()
    reader_dir = cache_root / f"{source.stem}-{digest[:12]}"
    manifest_path = reader_dir / "manifest.json"
    if manifest_path.exists() and not args.force:
        try:
            existing = json.loads(manifest_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            existing = {}
        if existing.get("tool_version") == TOOL_VERSION:
            print(json.dumps({"status": "reused", "reader_dir": str(reader_dir), "manifest": str(manifest_path)}, ensure_ascii=False, indent=2))
            return 0

    reader_dir.mkdir(parents=True, exist_ok=True)
    extraction_source = source
    method = "pdf-text"
    warnings: list[str] = []
    if source.suffix.lower() == ".docx":
        converted = convert_docx_to_pdf(source, reader_dir)
        if converted:
            extraction_source = converted
            method = "docx-via-libreoffice-pdf"
            pages, warnings = extract_pdf(extraction_source)
        else:
            method = "docx-structural-fallback"
            pages, warnings = extract_docx_fallback(source)
    else:
        pages, warnings = extract_pdf(source)

    chunks = make_chunks(pages, args.chunk_chars)
    chunks_dir = reader_dir / "chunks"
    if chunks_dir.exists():
        shutil.rmtree(chunks_dir)
    chunks_dir.mkdir(exist_ok=True)
    document_md = "\n\n".join(page_markdown(page) for page in pages)
    (reader_dir / "document.md").write_text(document_md, encoding="utf-8")
    for chunk in chunks:
        (reader_dir / chunk["file"]).write_text(chunk["text"], encoding="utf-8")
    index = build_index(chunks)
    (reader_dir / "index.json").write_text(json.dumps(index, ensure_ascii=False), encoding="utf-8")

    review_pages = [
        {"page": page.get("page"), "flags": page.get("flags", [])}
        for page in pages
        if page.get("flags")
    ]
    manifest = {
        "schema_version": 1,
        "tool_version": TOOL_VERSION,
        "source": str(source),
        "source_sha256": digest,
        "source_size_bytes": source.stat().st_size,
        "source_type": source.suffix.lower().lstrip("."),
        "extraction_method": method,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "reader_dir": str(reader_dir),
        "extraction_source": str(extraction_source),
        "page_count": len(pages) if pages and pages[0].get("page") is not None else None,
        "chunk_count": len(chunks),
        "chunk_target_chars": args.chunk_chars,
        "character_count": len(document_md),
        "warnings": warnings,
        "quality": {
            "text_pages": sum(1 for page in pages if len(page.get("text", "")) >= 8),
            "no_text_pages": sum(1 for page in pages if "no_text" in page.get("flags", [])),
            "flagged_pages": len(review_pages),
        },
        "visual_review_candidates": review_pages,
        "chunks": [{key: chunk[key] for key in ("id", "file", "page_start", "page_end", "chars", "flags")} for chunk in chunks],
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"status": "created", "reader_dir": str(reader_dir), "manifest": str(manifest_path), "summary": {
        "pages": manifest["page_count"], "chunks": len(chunks), "characters": len(document_md), "warnings": len(warnings), "visual_review_candidates": len(review_pages)
    }}, ensure_ascii=False, indent=2))
    return 0


def query(args: argparse.Namespace) -> int:
    reader_dir = Path(args.reader_dir).expanduser().resolve()
    index = json.loads((reader_dir / "index.json").read_text(encoding="utf-8"))
    query_terms = tokenize(args.question)
    if not query_terms:
        raise ValueError("Question contains no searchable terms")
    total = max(1, len(index))
    doc_frequency = Counter()
    for record in index:
        for term in set(record["terms"]):
            doc_frequency[term] += 1
    results = []
    normalized_query = re.sub(r"\s+", "", args.question.lower())
    for record in index:
        score = 0.0
        for term in query_terms:
            tf = record["terms"].get(term, 0)
            if tf:
                score += (1.0 + math.log(tf)) * (math.log((total + 1) / (doc_frequency[term] + 0.5)) + 1.0)
        chunk_path = reader_dir / record["file"]
        text = chunk_path.read_text(encoding="utf-8")
        if normalized_query and normalized_query in re.sub(r"\s+", "", text.lower()):
            score += 8.0
        if score > 0:
            snippet = make_snippet(text, args.question, query_terms)
            results.append({
                "score": round(score, 4), "id": record["id"], "file": str(chunk_path),
                "page_start": record["page_start"], "page_end": record["page_end"],
                "chars": record["chars"], "flags": record["flags"], "snippet": snippet,
            })
    results.sort(key=lambda item: (-item["score"], item["id"]))
    selected = results[: args.limit]
    print(json.dumps({"question": args.question, "reader_dir": str(reader_dir), "result_count": len(selected), "results": selected}, ensure_ascii=False, indent=2))
    return 0


def inspect_cache(args: argparse.Namespace) -> int:
    reader_dir = Path(args.reader_dir).expanduser().resolve()
    print((reader_dir / "manifest.json").read_text(encoding="utf-8"))
    return 0


def make_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--version", action="version", version=f"%(prog)s {TOOL_VERSION}")
    subparsers = parser.add_subparsers(dest="command", required=True)
    prep = subparsers.add_parser("prepare", help="Extract and index a PDF or DOCX")
    prep.add_argument("file")
    prep.add_argument("--output", help="Cache root; defaults to Documents/Codex/.document-reader-cache")
    prep.add_argument("--chunk-chars", type=int, default=DEFAULT_CHUNK_CHARS)
    prep.add_argument("--force", action="store_true")
    prep.set_defaults(func=prepare)
    search = subparsers.add_parser("query", help="Rank relevant chunks without printing full contents")
    search.add_argument("reader_dir")
    search.add_argument("question")
    search.add_argument("--limit", type=int, default=3)
    search.set_defaults(func=query)
    show = subparsers.add_parser("inspect", help="Print the compact manifest")
    show.add_argument("reader_dir")
    show.set_defaults(func=inspect_cache)
    return parser


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    parser = make_parser()
    args = parser.parse_args()
    try:
        return args.func(args)
    except Exception as exc:
        print(json.dumps({"error": str(exc), "type": type(exc).__name__}, ensure_ascii=False), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
